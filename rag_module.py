"""
RAG 모듈: 사업공고 문서를 FAISS 벡터 DB로 구축하고 검색/질의응답 수행
- PDF, HWP(한글) 파일 자동 감지 및 처리
- 개선된 질의응답 기능 (Top-k 검색 + 자연스러운 답변 생성)
- GPT-5 temperature=1 강제 설정 (Langchain 함수 이슈)
- LLM 기반 스마트 질문 생성 (도서 내용 이해하고 질문 생성)
"""
import os
import tempfile
from pathlib import Path
from typing import List, Optional, Tuple
import requests
import streamlit as st
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    UnstructuredFileLoader,
    TextLoader
)

class ProjectRAGSystem:
    """
    사업공고별 RAG 시스템
    - 다양한 문서 형식 다운로드 및 파싱 (PDF, HWP, DOCX, TXT)
    - FAISS 벡터 DB 구축
    - 개선된 질의응답 기능 (Top-k 검색 + 자연스러운 답변 생성)
    - LLM 기반 스마트 질문 생성 (도서 내용 이해하고 질문 생성)
    """
    
    def __init__(self, openai_api_key: str, llm_base_url: str, embedding_base_url: str = None):
        """
        Args:
            openai_api_key: OpenAI API 키
            llm_base_url: LLM 엔드포인트 URL
            embedding_base_url: Embeddings 엔드포인트 URL
        """
        print("\n" + "=" * 60)
        print("   ProjectRAGSystem 초기화 중...")
        print(f"   LLM URL: {llm_base_url}")
        print(f"   Embedding URL (전달받음): {embedding_base_url}")
        
        self.api_key = openai_api_key
        self.llm_base_url = llm_base_url
        
        # embedding_base_url 최종 설정
        if embedding_base_url and embedding_base_url.strip() and embedding_base_url != llm_base_url:
            self.embedding_base_url = embedding_base_url
            print(f"✅ Embedding URL 사용: {self.embedding_base_url}")
        else:
            self.embedding_base_url = llm_base_url
            print(f"❌ Embedding URL을 LLM URL로 대체: {self.embedding_base_url}")
        
        print("=" * 60 + "\n")
        
        # Embeddings 초기화 시도
        embedding_configs = [
            ("text-embedding-3-small", "base_url"),
            ("text-embedding-ada-002", "base_url"),
            ("openai/text-embedding-3-small", "base_url"),
            ("text-embedding-3-small", "openai_api_base"),
            ("openai/text-embedding-3-small", "openai_api_base"),
        ]
        
        self.embeddings = None
        last_error = None
        
        for model, param_type in embedding_configs:
            try:
                print(f"   시도: model={model}, {param_type}={self.embedding_base_url}")
                
                if param_type == "base_url":
                    self.embeddings = OpenAIEmbeddings(
                        openai_api_key=openai_api_key,
                        base_url=self.embedding_base_url,
                        model=model
                    )
                else:  # openai_api_base
                    self.embeddings = OpenAIEmbeddings(
                        openai_api_key=openai_api_key,
                        openai_api_base=self.embedding_base_url,
                        model=model
                    )
                
                # 테스트 임베딩
                test_result = self.embeddings.embed_query("테스트")
                print(f"✅ 성공! model={model}, 차원={len(test_result)}")
                break
                
            except Exception as e:
                last_error = str(e)
                print(f"❌ 실패: {last_error[:80]}")
                self.embeddings = None
                continue
        
        if not self.embeddings:
            error_msg = (
                f"\n{'='*60}\n"
                f"❌ Embeddings 모델 로드 실패\n"
                f"   마지막 오류: {last_error}\n"
                f"   Embedding URL: {self.embedding_base_url}\n"
                f"   API Key 앞 20자: {openai_api_key[:20]}...\n"
                f"{'='*60}\n"
            )
            print(error_msg)
            raise ValueError(error_msg)
        
        # LLM 초기화 (GPT-5는 temperature=1만 지원!)
        print(f"   LLM 초기화 중... (URL: {llm_base_url})")
        self.llm = ChatOpenAI(
            model="openai/gpt-5",
            openai_api_key=openai_api_key,
            base_url=llm_base_url,
            temperature=1,  # GPT-5 필수 설정!
        )
        print(f"✅ LLM 초기화 완료 (temperature=1 강제 설정)\n")
        
        self.vectorstore: Optional[FAISS] = None
    
    def detect_file_type(self, url: str, content: bytes) -> str:
        """
        URL과 파일 내용을 기반으로 파일 타입 감지
        
        Args:
            url: 파일 URL
            content: 파일 바이너리 내용
            
        Returns:
            파일 타입 ('pdf', 'hwp', 'docx', 'txt', 'unknown')
        """
        url_lower = url.lower()
        
        # HWP 파일 감지 (확장자 또는 매직 넘버)
        if url_lower.endswith('.hwp') or url_lower.endswith('.hwpx'):
            return 'hwp'
        
        # PDF 파일 감지
        if url_lower.endswith('.pdf') or content.startswith(b'%PDF'):
            return 'pdf'
        
        # DOCX 파일 감지
        if url_lower.endswith('.docx') or url_lower.endswith('.doc'):
            return 'docx'
        
        # TXT 파일
        if url_lower.endswith('.txt'):
            return 'txt'
        
        # 매직 넘버로 HWP 확인 (OLE 구조)
        if content.startswith(b'HWP Document File') or content.startswith(b'\xd0\xcf\x11\xe0'):
            return 'hwp'
        
        # 매직 넘버로 PDF 확인
        if content.startswith(b'%PDF'):
            return 'pdf'
        
        # DOCX는 ZIP 포맷 (PK로 시작)
        if content.startswith(b'PK\x03\x04'):
            return 'docx'
        
        return 'unknown'
    
    def download_file(self, url: str, save_path: str) -> Tuple[bool, str]:
        """
        URL에서 파일 다운로드 및 타입 감지
        
        Args:
            url: 파일 다운로드 URL
            save_path: 저장할 파일 경로 (확장자 제외)
            
        Returns:
            (성공 여부, 실제 저장된 파일 경로)
        """
        try:
            response = requests.get(url, timeout=30, verify=False)
            response.raise_for_status()
            
            content = response.content
            
            # 파일 타입 감지
            file_type = self.detect_file_type(url, content)
            
            # 적절한 확장자로 저장
            if file_type == 'pdf':
                final_path = f"{save_path}.pdf"
            elif file_type == 'hwp':
                final_path = f"{save_path}.hwp"
            elif file_type == 'docx':
                final_path = f"{save_path}.docx"
            elif file_type == 'txt':
                final_path = f"{save_path}.txt"
            else:
                # 알 수 없는 경우 원본 URL에서 확장자 추출 시도
                ext = Path(url).suffix or '.bin'
                final_path = f"{save_path}{ext}"
            
            # 파일 저장
            with open(final_path, 'wb') as f:
                f.write(content)
            
            return True, final_path
            
        except Exception as e:
            st.error(f"❌ 파일 다운로드 실패: {url}\n오류: {str(e)}")
            return False, ""
    
    def load_document_with_loader(self, file_path: str, url: str, doc_index: int) -> List[Document]:
        """
        파일 타입에 맞는 LangChain 로더를 사용하여 문서 로드
        
        Args:
            file_path: 로컬 파일 경로
            url: 원본 URL
            doc_index: 문서 인덱스
            
        Returns:
            Document 객체 리스트
        """
        file_ext = Path(file_path).suffix.lower()
        
        try:
            # PDF 파일
            if file_ext == '.pdf':
                loader = PyPDFLoader(file_path)
                docs = loader.load()
                st.success(f"✅ PDF 문서 로드 완료: {len(docs)} 페이지")
            
            # HWP 파일 - Windows 특화 처리
            elif file_ext in ['.hwp', '.hwpx']:
                st.info(f"🔄 HWP 파일 처리 중... (시간이 걸릴 수 있습니다)")
                
                try:
                    # olefile로 텍스트 추출
                    import olefile
                    
                    if not olefile.isOleFile(file_path):
                        st.error("❌ 올바른 HWP 파일이 아닙니다.")
                        return []
                    
                    ole = olefile.OleFileIO(file_path)
                    
                    # HWP 파일에서 텍스트 스트림 찾기
                    text_content = ""
                    
                    # 방법 1: PrvText 스트림 (미리보기 텍스트)
                    if ole.exists('PrvText'):
                        encoded_text = ole.openstream('PrvText').read()
                        
                        # 여러 인코딩 시도
                        for encoding in ['utf-16', 'utf-16-le', 'cp949', 'euc-kr']:
                            try:
                                text_content = encoded_text.decode(encoding, errors='ignore')
                                if text_content.strip():
                                    break
                            except:
                                continue
                    
                    # 방법 2: BodyText 스트림 시도
                    if not text_content.strip() and ole.exists('BodyText'):
                        try:
                            encoded_text = ole.openstream('BodyText').read()
                            text_content = encoded_text.decode('utf-16', errors='ignore')
                        except:
                            pass
                    
                    ole.close()
                    
                    if text_content.strip():
                        # Document 객체 생성
                        doc = Document(
                            page_content=text_content,
                            metadata={'source': file_path}
                        )
                        docs = [doc]
                        st.success(f"✅ HWP 텍스트 추출 완료 ({len(text_content)} 글자)")
                    else:
                        st.error("❌ HWP 파일에서 텍스트를 추출할 수 없습니다.")
                        st.info("   이 HWP 파일을 PDF로 변환 후 다시 시도해주세요.")
                        return []
                        
                except Exception as e:
                    st.error(f"❌ HWP 처리 실패: {str(e)}")
                    st.info("   HWP 파일을 PDF로 변환하거나, 다른 문서를 사용해주세요.")
                    return []
            
            # DOCX 파일
            elif file_ext in ['.docx', '.doc']:
                try:
                    # python-docx 직접 사용
                    from docx import Document as DocxDocument
                    
                    docx = DocxDocument(file_path)
                    text_content = "\n".join([para.text for para in docx.paragraphs if para.text.strip()])
                    
                    if text_content.strip():
                        doc = Document(
                            page_content=text_content,
                            metadata={'source': file_path}
                        )
                        docs = [doc]
                        st.success(f"✅ DOCX 문서 로드 완료 ({len(text_content)} 글자)")
                    else:
                        st.warning("⚠️ DOCX 파일이 비어있습니다. 건너뜁니다.")
                        return []
                except Exception as e:
                    st.warning(f"❌ DOCX 처리 실패, 건너뜁니다: {str(e)}")
                    return []
            
            # TXT 파일
            elif file_ext == '.txt':
                encodings = ['utf-8', 'cp949', 'euc-kr', 'utf-8-sig']
                docs = None
                for encoding in encodings:
                    try:
                        loader = TextLoader(file_path, encoding=encoding)
                        docs = loader.load()
                        break
                    except:
                        continue
                
                if docs:
                    st.success(f"✅ TXT 문서 로드 완료")
                else:
                    st.error("❌ TXT 파일 인코딩 감지 실패")
                    return []
            
            # 기타 파일
            else:
                try:
                    loader = UnstructuredFileLoader(file_path, mode="single")
                    docs = loader.load()
                    st.success(f"✅ 문서 로드 완료 ({file_ext})")
                except:
                    st.error(f"❌ {file_ext} 파일 형식은 지원하지 않습니다.")
                    return []
            
            # 메타데이터 추가
            for doc in docs:
                doc.metadata['source_url'] = url
                doc.metadata['doc_index'] = doc_index
                doc.metadata['file_type'] = file_ext
            
            return docs
            
        except Exception as e:
            st.error(f"❌ 문서 로드 실패: {file_path}\n오류: {str(e)}")
            return []
    
    def load_documents_from_urls(self, file_urls: List[str]) -> List[Document]:
        all_documents = []
        temp_dir = tempfile.mkdtemp()
        
        # 빈 URL 필터링 강화
        valid_urls = []
        for url in file_urls:
            if url and isinstance(url, str) and url.strip():
                url_clean = url.strip()
                if url_clean.lower() not in ['nan', 'none', 'null', '']:
                    valid_urls.append(url_clean)
        
        if not valid_urls:
            st.warning("⚠️ 유효한 문서 URL이 없습니다.")
            return []
        
        st.info(f"   총 {len(valid_urls)}개의 문서를 다운로드합니다...")
        
        success_count = 0
        fail_count = 0
        
        for idx, url in enumerate(valid_urls):
            try:
                st.info(f"   문서 {idx+1}/{len(valid_urls)} 처리 중...")
                
                temp_file_path = os.path.join(temp_dir, f"doc_{idx}")
                
                # 파일 다운로드
                success, actual_path = self.download_file(url, temp_file_path)
                
                if success and actual_path:
                    # 문서 로드
                    docs = self.load_document_with_loader(actual_path, url, idx)
                    
                    if docs and len(docs) > 0:
                        # 빈 문서 체크
                        valid_docs = [d for d in docs if d.page_content.strip()]
                        if valid_docs:
                            all_documents.extend(valid_docs)
                            success_count += 1
                        else:
                            st.warning(f"⚠️ 문서 {idx+1}: 내용이 비어있어 건너뜁니다.")
                            fail_count += 1
                    else:
                        fail_count += 1
                    
                    # 임시 파일 삭제
                    try:
                        if os.path.exists(actual_path):
                            os.remove(actual_path)
                    except:
                        pass
                else:
                    fail_count += 1
            except Exception as e:
                st.error(f"❌ 문서 {idx+1} 처리 중 오류: {str(e)}")
                fail_count += 1
                continue
        
        # 결과 요약
        if all_documents:
            st.success(f"🎉 성공: {success_count}개 문서, 총 {len(all_documents)}개 청크 로드!")
            if fail_count > 0:
                st.warning(f"⚠️ 건너뜀/실패: {fail_count}개 문서")
        else:
            st.error("❌ 로드된 문서가 없습니다.")
        
        return all_documents
    
    def build_vectorstore(self, documents: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200):
        if not documents:
            st.error("❌ 로드된 문서가 없습니다.")
            return
        
        # 텍스트 분할
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        
        splits = text_splitter.split_documents(documents)
        st.info(f"   총 {len(splits)}개의 텍스트 청크 생성됨")
        
        # 디버깅: 설정 확인
        st.info(f"🔍 Embedding Base URL: {self.embedding_base_url}")
        st.info(f"🔍 API Key 앞 10자: {self.api_key[:5]}...")
        
        try:
            # 테스트 임베딩
            st.info("🧪 임베딩 테스트 중...")
            test_embedding = self.embeddings.embed_query("테스트")
            st.success(f"✅ 임베딩 테스트 성공! (차원: {len(test_embedding)})")
            
            # 벡터스토어 생성
            st.info("🔄 벡터스토어 생성 중...")
            self.vectorstore = FAISS.from_documents(splits, self.embeddings)
            
            st.success("✅ 벡터 DB 구축 완료!")
            
        except Exception as e:
            st.error(f"❌ 벡터 DB 구축 실패: {str(e)}")
            st.error(f"🔍 사용 중인 Embedding URL: {self.embedding_base_url}")
            
            # 상세 오류 출력
            import traceback
            with st.expander("상세 오류 보기"):
                st.code(traceback.format_exc())
    
    # ⭐ 신규 메서드: LLM 기반 스마트 질문 생성
    def generate_smart_questions(self, num_questions: int = 4) -> List[str]:
        """
        LLM을 사용해 문서 내용 기반 질문 자동 생성
        
        Args:
            num_questions: 생성할 질문 개수 (기본 4개)
            
        Returns:
            생성된 질문 리스트
        """
        if not self.vectorstore:
            return [
                "이 사업의 주요 내용을 요약해주세요.",
                "지원 대상은 누구인가요?",
            ]
        
        try:
            # 1. 문서 샘플 추출 (다양한 부분에서)
            sample_docs = self.vectorstore.similarity_search("", k=5)
            
            # 2. 샘플 텍스트 구성 (각 문서 앞부분 500자)
            sample_texts = []
            for i, doc in enumerate(sample_docs[:3], 1):  # 최대 3개 문서
                content_preview = doc.page_content[:500].strip()
                sample_texts.append(f"[문서 {i}]\n{content_preview}")
            
            combined_sample = "\n\n".join(sample_texts)
            
            # 3. LLM 프롬프트 생성
            prompt = f"""당신은 정부지원사업 공고 문서 분석 전문가입니다.
아래 문서 일부를 읽고, 사용자가 이 문서에 대해 가장 궁금해할 만한 질문 {num_questions}개를 생성하세요.

[문서 내용 샘플]
{combined_sample}

[중요 지침]
1. 문서에 답이 명확히 존재하는 질문만 생성하세요
2. 각 질문은 한 문장으로 간결하게 작성하세요
3. 실용적이고 구체적인 질문이어야 합니다
4. 다음 주제를 우선적으로 다루세요:
   - 지원 대상 및 자격
   - 신청 방법 및 절차
   - 지원 내용 및 규모
   - 제출 서류
   - 평가 기준
   - 사업 기간

[출력 형식]
1. 질문1
2. 질문2
3. 질문3
4. 질문4

위 형식으로만 출력하세요. 추가 설명 없이 질문만 나열하세요."""

            # 4. LLM 호출
            response = self.llm.invoke(prompt)
            questions_text = response.content if hasattr(response, "content") else str(response)
            
            # 5. 파싱
            questions = []
            for line in questions_text.split("\n"):
                line = line.strip()
                # "1. " 또는 "- " 형식 파싱
                if line and (line[0].isdigit() or line.startswith("-")):
                    # 번호/기호 제거
                    if "." in line:
                        clean_q = line.split(".", 1)[-1].strip()
                    else:
                        clean_q = line.lstrip("- ").strip()
                    
                    if clean_q and len(clean_q) > 5:  # 최소 길이 체크
                        questions.append(clean_q)
            
            # 6. 결과 검증
            if len(questions) >= 2:
                return questions[:num_questions]
            else:
                # Fallback: 기본 질문
                return [
                    "이 사업의 주요 목적과 내용을 요약해주세요.",
                    "지원 대상과 신청 자격은 무엇인가요?",
                    "어떤 방식으로 지원을 받을 수 있나요?",
                    "신청 시 제출해야 할 서류는 무엇인가요?",
                ][:num_questions]
                
        except Exception as e:
            print(f"⚠️ 질문 생성 실패: {e}")
            # Fallback: 기본 질문
            return [
                "이 사업의 핵심 내용을 설명해주세요.",
                "누가 이 사업에 지원할 수 있나요?",
                "지원 절차는 어떻게 되나요?",
                "어떤 혜택을 받을 수 있나요?",
            ][:num_questions]
    
    def query(self, question: str, use_enhanced_prompt: bool = True) -> Tuple[str, List[Document]]:
        """
        질문에 대한 답변 생성 (개선된 버전)
        
        Args:
            question: 사용자 질문
            use_enhanced_prompt: 향상된 프롬프트 사용 여부
            
        Returns:
            (답변, 참조 문서 리스트)
        """
        if not self.vectorstore:
            return "❌ 벡터 DB가 구축되지 않았습니다. 먼저 문서를 로드해주세요.", []
        
        try:
            # 1. 관련 문서 검색 (Top-5)
            source_docs = self.vectorstore.similarity_search(question, k=5)
            
            if not source_docs:
                return "❌ 관련 문서를 찾을 수 없습니다. 다른 질문을 시도해보세요.", []
            
            # 2. 컨텍스트 구성
            context = "\n\n---\n\n".join([
                f"[문서 {i+1}]\n{doc.page_content}" 
                for i, doc in enumerate(source_docs)
            ])
            
            # 3. 프롬프트 생성
            if use_enhanced_prompt:
                prompt = f"""당신은 정부지원사업 공고 문서 분석 전문가입니다. 
아래 문서 내용을 바탕으로 사용자의 질문에 정확하고 친절하게 답변해주세요.

[주요 지침]
1. 문서에 명시된 내용만을 기반으로 답변하세요
2. 정확한 수치, 날짜, 조건 등은 원문 그대로 인용하세요
3. 문서에 없는 내용은 "문서에 명시되어 있지 않습니다"라고 답하세요
4. 답변은 구조화되고 읽기 쉽게 작성하세요 (필요시 글머리기호 사용)
5. 전문용어는 쉽게 풀어서 설명해주세요
6. 답변은 한국어로 자연스럽게 작성하세요

[참조 문서 내용]
{context}

[사용자 질문]
{question}

[답변]"""
            else:
                prompt = f"""다음 문서 내용을 참고하여 질문에 답변해주세요.

문서 내용:
{context}

질문: {question}

답변:"""
            
            # 4. LLM 호출 (temperature는 이미 초기화 시 설정됨)
            response = self.llm.invoke(prompt)
            answer = response.content if hasattr(response, "content") else str(response)
            
            return answer, source_docs
            
        except Exception as e:
            import traceback
            error_detail = traceback.format_exc()
            return f"❌ 오류 발생: {str(e)}\n\n{error_detail}", []
    
    def similarity_search(self, query: str, k: int = 5) -> List[Tuple[Document, float]]:
        """
        유사도 기반 문서 검색
        
        Args:
            query: 검색 쿼리
            k: 반환할 문서 개수
            
        Returns:
            (Document, 유사도 점수) 리스트
        """
        if not self.vectorstore:
            return []
        
        try:
            results = self.vectorstore.similarity_search_with_score(query, k=k)
            return results
        except Exception as e:
            st.error(f"❌ 검색 오류: {str(e)}")
            return []


def extract_file_urls(project_row) -> List[str]:
    """
    프로젝트 row에서 공고규격서URL1~10 추출
    
    Args:
        project_row: 사업공고 DataFrame row
        
    Returns:
        유효한 파일 URL 리스트 (PDF, HWP 등)
    """
    file_urls = []
    
    for i in range(1, 11):
        col_name = f"공고규격서URL{i}"
        if col_name in project_row.index:
            url = project_row[col_name]
            if url and str(url).strip() and str(url) != 'nan':
                file_urls.append(str(url).strip())
    
    return file_urls


@st.cache_resource
def get_rag_system(api_key: str, llm_base_url: str, embedding_base_url: str = None):
    """
    RAG 시스템 인스턴스 생성 (캐싱)
    """
    return ProjectRAGSystem(api_key, llm_base_url, embedding_base_url)


def display_source_documents(source_docs: List[Document]):
    """
    참조 문서 표시 (Streamlit UI)
    
    Args:
        source_docs: 참조된 Document 리스트
    """
    if not source_docs:
        return
    
    st.markdown("### 📚 참조 문서")
    
    for idx, doc in enumerate(source_docs, 1):
        with st.expander(f"참조 {idx}: {doc.metadata.get('source_url', 'Unknown')[:250]}..."):
            st.markdown(f"**출처:** {doc.metadata.get('source_url', 'N/A')}")
            st.markdown(f"**파일 타입:** {doc.metadata.get('file_type', 'N/A')}")
            st.markdown(f"**페이지:** {doc.metadata.get('page', 'N/A')}")
            st.markdown("**내용:**")
            content = doc.page_content[:500] + "..." if len(doc.page_content) > 500 else doc.page_content
            st.text(content)